"""
Fill reimbursement forms using the template file.

This script copies the template reimbursement form to each trip folder
and fills in the trip information based on the Zhejiang University form structure.
"""

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt
import sys
import os

# Try to import win32com for .doc file support
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False
    print("Warning: win32com not available, cannot process .doc files")
    print("Install with: pip install pywin32")


@dataclass
class InvoiceInfo:
    """Invoice information extracted from filename."""
    date: datetime
    invoice_type: str  # 机票, 火车, 接送机, 打车, 住宿, etc.
    origin: Optional[str]
    destination: Optional[str]
    amount: float
    traveler: str = ""
    document_type: str = ""

    @classmethod
    def from_filename(cls, filename: str) -> 'InvoiceInfo':
        """Extract invoice info from filename.

        Supports multiple formats:
        - Flight/Train: {date}_{type}_{origin}_{destination}_{amount}_{traveler}_{doc_type}.pdf
          Example: 2026-02-10_机票_杭州_青岛_1065.00_王春晖.pdf
        - Airport Transfer: {date_range}_{type}_{route}_{amount}_{traveler}_{doc_type}.pdf
          Example: 2026-03-01至2026-03-01_接送机_首都国际机场_T3_北京华融大厦_133.00_王春晖_行程单.pdf
        - Taxi (simple): {date_range}_{type}_{amount}_{traveler}_{doc_type}.pdf
          Example: 2026-01-28至2026-01-28_打车_17.40_王春晖_发票.pdf
        - Accommodation: {date_range}_{type}_{amount}_{traveler}.pdf
          Example: 2026-02-24_2026-02-25_住宿_342.02_王春晖.pdf
        """
        name = filename.replace('.pdf', '')
        parts = name.split('_')

        if len(parts) < 4:
            return None

        try:
            # Parse date - handle single date and date range formats
            date_str = parts[0]
            if '至' in date_str:
                # Date range format - extract start date
                start_date_str = date_str.split('至')[0]
                invoice_date = datetime.strptime(start_date_str, "%Y-%m-%d")
            elif len(parts) > 1 and re.match(r'\d{4}-\d{2}-\d{2}', parts[1]):
                # Accommodation format: 2026-02-24_2026-02-25
                invoice_date = datetime.strptime(parts[0], "%Y-%m-%d")
            else:
                # Single date format
                invoice_date = datetime.strptime(date_str, "%Y-%m-%d")

            # Find invoice type
            invoice_type = None
            type_idx = 1
            for i, part in enumerate(parts[1:], start=1):
                if part in ["机票", "火车", "接送机", "打车", "住宿", "餐饮", "结账单"]:
                    invoice_type = part
                    type_idx = i
                    break

            if not invoice_type:
                if len(parts) > 1:
                    invoice_type = parts[1]
                else:
                    return None

            # Parse origin, destination, amount, traveler
            origin = None
            destination = None
            amount = 0.0
            traveler = ""
            doc_type = ""

            # Parse based on invoice type
            if invoice_type in ["机票", "火车"]:
                # Format: {date}_{type}_{origin}_{destination}_{amount}_{traveler}_{doc_type}
                # OR: {date}_{type}_{origin}_{amount}_{traveler}
                if len(parts) >= 5:
                    origin = parts[2]
                    if re.match(r'^\d+\.\d{2}$', parts[3]):
                        amount = float(parts[3])
                        if len(parts) > 4:
                            traveler = parts[4]
                        if len(parts) > 5 and parts[5] in ["行程单", "发票"]:
                            doc_type = parts[5]
                    else:
                        destination = parts[3]
                        for i in range(3, len(parts)):
                            amount_match = re.match(r'(\d+\.\d{2})', parts[i])
                            if amount_match:
                                amount = float(amount_match.group(1))
                                if i + 1 < len(parts):
                                    traveler = parts[i + 1]
                                if i + 2 < len(parts) and parts[i + 2] in ["行程单", "发票"]:
                                    doc_type = parts[i + 2]
                                break

            elif invoice_type == "接送机":
                # Format: {date}_{type}_{route_parts}_{amount}_{traveler}_{doc_type}
                # OR: {date}_{type}_{amount}_{traveler}_{doc_type}
                amount_idx = None
                for i in range(2, len(parts)):
                    if re.match(r'^\d+\.\d{2}$', parts[i]):
                        amount_idx = i
                        break

                if amount_idx and amount_idx > 2:
                    origin = "_".join(parts[2:amount_idx])
                    amount = float(parts[amount_idx])
                    if amount_idx + 1 < len(parts):
                        traveler = parts[amount_idx + 1]
                    if amount_idx + 2 < len(parts) and parts[amount_idx + 2] in ["行程单", "发票"]:
                        doc_type = parts[amount_idx + 2]
                elif amount_idx:
                    amount = float(parts[amount_idx])
                    if amount_idx + 1 < len(parts):
                        traveler = parts[amount_idx + 1]
                    if amount_idx + 2 < len(parts) and parts[amount_idx + 2] in ["行程单", "发票"]:
                        doc_type = parts[amount_idx + 2]

            elif invoice_type in ["打车", "餐饮"]:
                # Format: {date_range}_{type}_{amount}_{traveler}_{doc_type}
                for i in range(2, len(parts)):
                    amount_match = re.match(r'(\d+\.\d{2})', parts[i])
                    if amount_match:
                        amount = float(amount_match.group(1))
                        if i + 1 < len(parts):
                            traveler = parts[i + 1]
                        if i + 2 < len(parts) and parts[i + 2] in ["行程单", "发票"]:
                            doc_type = parts[i + 2]
                        break

            elif invoice_type in ["住宿", "结账单"]:
                # Format: {date_range}_{type}_{amount}_{traveler}
                for i in range(2, len(parts)):
                    amount_match = re.match(r'(\d+\.\d{2})', parts[i])
                    if amount_match:
                        amount = float(amount_match.group(1))
                        if i + 1 < len(parts):
                            traveler = parts[i + 1]
                        break
                # Normalize type
                if invoice_type == "结账单":
                    invoice_type = "住宿"

            return cls(
                date=invoice_date,
                invoice_type=invoice_type,
                origin=origin,
                destination=destination,
                amount=amount,
                traveler=traveler,
                document_type=doc_type
            )
        except Exception as e:
            print(f"Error parsing {filename}: {e}")
            return None


def read_trip_invoices(trip_dir: Path) -> List[InvoiceInfo]:
    """Read all invoice files from a trip directory."""
    invoices = []
    for pdf_file in trip_dir.glob("*.pdf"):
        inv = InvoiceInfo.from_filename(pdf_file.name)
        if inv:
            invoices.append(inv)

    # Remove duplicates: keep only "发票" version when both "发票" and "行程单" exist
    # Group by (date, type, amount) and keep invoice over itinerary
    unique_invoices = {}
    for inv in invoices:
        # Create key: (date, type, amount) - document_type is the only difference
        key = (inv.date, inv.invoice_type, inv.amount)
        if key not in unique_invoices:
            unique_invoices[key] = inv
        else:
            # Prefer "发票" over "行程单"
            existing = unique_invoices[key]
            if (inv.document_type == "发票" and existing.document_type != "发票") or \
               (inv.document_type == "" and existing.document_type == "行程单"):
                unique_invoices[key] = inv
            # Otherwise keep existing

    return sorted(unique_invoices.values(), key=lambda x: x.date)


def parse_trip_folder(folder_name: str) -> tuple:
    """Parse trip folder name to extract destination and dates.

    Format: YYYYMMDD_YYYYMMDD_Destination
    Returns: (start_date_str, end_date_str, destination)
    """
    parts = folder_name.split('_')
    if len(parts) >= 3:
        start_date = parts[0]  # YYYYMMDD
        end_date = parts[1]    # YYYYMMDD
        destination = parts[2].replace('-', '、')
        return start_date, end_date, destination
    return "", "", ""


def format_date_range(start_date_str: str, end_date_str: str) -> str:
    """Convert YYYYMMDD-YYYYMMDD to MM月DD日-MM月DD日 format."""
    def fmt(d):
        return f"{d[4:6]}月{d[6:8]}日"
    return f"{fmt(start_date_str)}-{fmt(end_date_str)}"


def number_to_chinese(num: float) -> str:
    """Convert number to Chinese uppercase (for amount in words)."""
    units = ['', '拾', '佰', '仟', '万']
    digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']

    integer_part = int(num)
    decimal_part = round((num - integer_part) * 100)

    result = ""
    if integer_part == 0:
        result = "零"
    else:
        str_num = str(integer_part)
        for i, digit in enumerate(str_num):
            result += digits[int(digit)]
            pos = len(str_num) - i - 1
            if pos > 0:
                result += units[pos] if digit != '0' else ''

    result += "元"
    if decimal_part > 0:
        result += f"{digits[decimal_part // 10]}角{digits[decimal_part % 10]}分"
    else:
        result += "整"

    return result


def fill_cell_text(cell, text: str):
    """Helper to set cell text, preserving formatting."""
    for para in cell.paragraphs:
        para.clear()

    if not cell.paragraphs:
        cell.add_paragraph()
    cell.paragraphs[0].text = text
    cell.paragraphs[0].runs[0].font.size = Pt(10)


def convert_doc_to_docx(doc_path: Path) -> Path:
    """Convert .doc file to .docx using win32com."""
    if not HAS_WIN32COM:
        raise RuntimeError("win32com is required for .doc to .docx conversion")

    docx_path = doc_path.with_suffix('.docx')

    # Skip if already converted
    if docx_path.exists():
        return docx_path

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        doc = word.Documents.Open(str(doc_path.absolute()))
        # FileFormat 16 represents wdFormatXMLDocument (.docx)
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)
        doc.Close()
        print(f"Converted: {doc_path} -> {docx_path}")
        return docx_path
    except Exception as e:
        print(f"Error converting {doc_path}: {e}")
        raise
    finally:
        word.Quit()


def fill_template_with_trip_info(template_path: Path, trip_dir: Path, invoices: List[InvoiceInfo]):
    """Copy template to trip folder and fill in trip information.

    Based on Zhejiang University reimbursement form structure:
    - Row 0, Col 3/4/5: 姓名
    - Row 1, Col 3/4/5: 职称
    - Row 2, Col 3/4/5: 出差事由
    - Row 6, Col 0: 出差地点
    - Row 6, Col 1: 起止日期
    - Row 6, Col 2: 天数
    - Row 6, Col 3: 飞机
    - Row 6, Col 5: 火车
    - Row 6, Col 6: 市内交通费 (打车、接送机)
    - Row 6, Col 7/8: 住宿费
    - Row 11, Col 12: 合计金额
    """
    # Determine output format based on template
    if template_path.suffix == '.doc':
        output_path = trip_dir / '差旅费报销单.doc'
        # Convert .doc to .docx for editing
        temp_docx = convert_doc_to_docx(template_path)
        edit_path = temp_docx
        save_as_doc = True
    else:
        output_path = trip_dir / '差旅费报销单.docx'
        edit_path = template_path
        save_as_doc = False

    if output_path.exists():
        output_path.unlink()

    shutil.copy2(template_path, output_path)

    doc = Document(edit_path)

    start_date, end_date, destination = parse_trip_folder(trip_dir.name)
    date_range = format_date_range(start_date, end_date)

    # Calculate totals by transport type
    airplane_total = 0.0
    train_total = 0.0
    local_transport_total = 0.0  # 打车、接送机 -> 市内交通费
    accommodation_total = 0.0

    for inv in invoices:
        if inv.invoice_type == '机票':
            airplane_total += inv.amount
        elif inv.invoice_type == '火车':
            train_total += inv.amount
        elif inv.invoice_type in ['接送机', '打车']:
            # 市内交通费
            local_transport_total += inv.amount
        elif inv.invoice_type == '住宿':
            accommodation_total += inv.amount

    total = airplane_total + train_total + local_transport_total + accommodation_total

    table = doc.tables[0]

    # Row 0: 姓名
    fill_cell_text(table.rows[0].cells[3], "王春晖")

    # Row 1: 职称
    fill_cell_text(table.rows[1].cells[3], "副研究员")

    # Row 2: 出差事由
    fill_cell_text(table.rows[2].cells[3], "会议、试验")

    # Row 6: First data row
    # Col 0: 出差地点
    fill_cell_text(table.rows[6].cells[0], destination)

    # Col 1: 起止日期
    fill_cell_text(table.rows[6].cells[1], date_range)

    # Col 2: 天数
    start = datetime.strptime(start_date, "%Y%m%d")
    end = datetime.strptime(end_date, "%Y%m%d")
    days = (end - start).days + 1
    fill_cell_text(table.rows[6].cells[2], str(days))

    # Col 3: 飞机
    if airplane_total > 0:
        fill_cell_text(table.rows[6].cells[3], f"{airplane_total:.2f}")

    # Col 5: 火车
    if train_total > 0:
        fill_cell_text(table.rows[6].cells[5], f"{train_total:.2f}")

    # Col 6: 市内交通费 (打车、接送机)
    if local_transport_total > 0:
        fill_cell_text(table.rows[6].cells[6], f"{local_transport_total:.2f}")

    # Col 7/8: 住宿费
    if accommodation_total > 0:
        fill_cell_text(table.rows[6].cells[7], f"{accommodation_total:.2f}")

    # Row 11: 合计金额
    fill_cell_text(table.rows[11].cells[12], f"{total:.2f}")

    # Save the modified document
    # If original was .doc, save as .docx (python-docx doesn't support .doc output)
    if save_as_doc:
        # Save as .docx and rename to .doc for consistency
        temp_output = trip_dir / '差旅费报销单_temp.docx'
        doc.save(temp_output)
        # Copy to final location
        shutil.copy2(temp_output, output_path)
        temp_output.unlink()
    else:
        doc.save(output_path)

    print(f"Created: {output_path}")

    return output_path


def generate_all_forms(template_path: str, trips_base_dir: str = "trips"):
    """Generate reimbursement forms for all trips using template."""
    template = Path(template_path)

    if not template.exists():
        print(f"Template file not found: {template_path}")
        return

    base_dir = Path(trips_base_dir)

    for traveler_dir in base_dir.iterdir():
        if not traveler_dir.is_dir() or traveler_dir.name.startswith('.'):
            continue

        print(f"\nProcessing traveler: {traveler_dir.name}")

        for trip_dir in traveler_dir.iterdir():
            if not trip_dir.is_dir():
                continue

            # Skip "普通打车" folder
            if trip_dir.name == "普通打车":
                continue

            print(f"  Processing trip: {trip_dir.name}")

            invoices = read_trip_invoices(trip_dir)

            if not invoices:
                print(f"    No invoices found, skipping")
                continue

            total = sum(inv.amount for inv in invoices)
            print(f"    Found {len(invoices)} invoices, total: {total:.2f} yuan")

            fill_template_with_trip_info(template, trip_dir, invoices)


if __name__ == "__main__":
    template_path = "trips/报销单.doc"
    generate_all_forms(template_path, "trips")
