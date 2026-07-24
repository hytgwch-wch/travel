"""生成 Word (.docx) 统计单。

每个行程 / 孤儿单据 / 市内交通 / 普通打车 文件夹生成一份「统计单.docx」，
附在该子文件夹下。不依赖外部模板，直接用 python-docx 从文件夹内发票构建。

入口：generate_for_folder(folder, kind) -> Path
  kind ∈ {"trip", "市内交通", "孤儿单据", "普通打车"}
"""
from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)

SUMMARY_FILENAME = "统计单.docx"
DEFAULT_FONT = "宋体"

# 文件夹名 → 统计单类型
SPECIAL_KINDS = ("市内交通", "孤儿单据", "普通打车")


# ---------------------------------------------------------------------------
# 发票信息（duck-typed：兼容 trip_grouper.Invoice，也兼容 fill_reimbursement 的对象）
# ---------------------------------------------------------------------------
def _route_str(inv) -> str:
    """线路/备注字符串。"""
    t = getattr(inv, "invoice_type", "") or ""
    if t in ("机票", "火车"):
        o = getattr(inv, "origin", None) or ""
        d = getattr(inv, "destination", None) or ""
        if o or d:
            return f"{o}→{d}".strip("→")
        return ""
    if t == "接送机":
        o = getattr(inv, "origin", None) or ""
        return o.replace("_", "→") if o else ""
    if getattr(inv, "is_refund", False):
        return "退票费"
    return ""


def _category_totals(invoices: list) -> dict:
    totals = {"机票": 0.0, "火车": 0.0, "市内交通": 0.0, "住宿": 0.0, "其他": 0.0}
    for inv in invoices:
        t = getattr(inv, "invoice_type", "") or ""
        if t == "机票":
            totals["机票"] += inv.amount
        elif t == "火车":
            totals["火车"] += inv.amount
        elif t in ("接送机", "打车"):
            totals["市内交通"] += inv.amount
        elif t == "住宿":
            totals["住宿"] += inv.amount
        else:
            totals["其他"] += inv.amount
    return totals


def _dedupe(invoices: list) -> list:
    """同一 (日期, 类型, 金额) 的 发票/行程单 只保留一份，优先「发票」。"""
    unique = {}
    for inv in invoices:
        key = (inv.date, inv.invoice_type, round(inv.amount, 2))
        if key not in unique:
            unique[key] = inv
            continue
        existing = unique[key]
        existing_dt = getattr(existing, "document_type", "") or ""
        inv_dt = getattr(inv, "document_type", "") or ""
        if (inv_dt == "发票" and existing_dt != "发票") or \
           (inv_dt == "" and existing_dt == "行程单"):
            unique[key] = inv
    return sorted(unique.values(), key=lambda x: (x.date, x.invoice_type))


def collect_folder_invoices(folder: Path) -> List:
    """扫描文件夹内 *.pdf → 解析 → 去重，返回统计口径的发票列表。

    口径：sorted glob → Invoice.from_filename → 丢弃解析失败/无 date 的
    → _dedupe（同 (日期, 类型, 金额) 的发票/行程单只保留一份，优先发票）。
    generate_for_folder（统计单 .docx）与 web 视图模型共用本函数，确保
    张数/金额口径一致，避免「发票+行程单」成对单据被各计一次。
    """
    from src.trip_grouper import Invoice  # 延迟导入，避免循环引用

    folder = Path(folder)
    invoices: List = []
    for pdf in sorted(folder.glob("*.pdf")):
        inv = Invoice.from_filename(pdf)
        if inv and getattr(inv, "date", None):
            invoices.append(inv)
    invoices = _dedupe(invoices)
    # 行程单（document_type=="行程单"）是服务记录，不计入金额统计——与 trip_grouper
    # 市内交通口径一致。含无配对发票的孤立行程单（如接送机费用已并入机票发票，
    # 仅有行程单无发票），避免被当作发票重复计费（[[bug-038]]）。
    return [inv for inv in invoices
            if (getattr(inv, "document_type", "") or "") != "行程单"]


def _number_to_chinese(num: float) -> str:
    """金额转中文大写。"""
    units = ['', '拾', '佰', '仟', '万']
    digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    integer_part = int(num)
    decimal_part = round((num - integer_part) * 100)
    if integer_part == 0:
        result = "零"
    else:
        result = ""
        s = str(integer_part)
        for i, ch in enumerate(s):
            d = int(ch)
            result += digits[d]
            pos = len(s) - i - 1
            if pos > 0 and d != 0:
                result += units[pos]
    result += "元"
    if decimal_part > 0:
        result += f"{digits[decimal_part // 10]}角{digits[decimal_part % 10]}分"
    else:
        result += "整"
    return result


# ---------------------------------------------------------------------------
# 元信息（从文件夹路径解析）
# ---------------------------------------------------------------------------
def _meta_from_folder(folder: Path, kind: str) -> dict:
    """folder = trips/{traveler}/{company}/{foldername} —— 父=company，祖父=traveler"""
    traveler = folder.parents[1].name if len(folder.parents) >= 2 else ""
    company = folder.parents[0].name if len(folder.parents) >= 1 else ""
    meta = {"traveler": traveler, "company": company, "kind": kind}
    if kind == "trip":
        name_parts = folder.name.split('_')
        if len(name_parts) >= 3:
            start, end = name_parts[0], name_parts[1]
            dest = name_parts[2]
            meta["start_date"] = f"{start[:4]}-{start[4:6]}-{start[6:8]}"
            meta["end_date"] = f"{end[:4]}-{end[4:6]}-{end[6:8]}"
            meta["destination"] = dest
        meta["title"] = "差旅费用统计单"
    elif kind == "市内交通":
        meta["title"] = "市内交通统计单"
        meta["subtitle"] = "接送机 / 打车（不计入单次出差）"
    elif kind == "孤儿单据":
        meta["title"] = "孤儿单据统计单"
        meta["subtitle"] = "退票费、独立住宿等（未形成完整行程）"
    else:  # 普通打车
        meta["title"] = "普通打车统计单"
        meta["subtitle"] = "未归入行程的打车票"
    return meta


# ---------------------------------------------------------------------------
# docx 构建
# ---------------------------------------------------------------------------
def _set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)


def _build_docx(folder: Path, invoices: list, meta: dict) -> Path:
    doc = Document()

    # Normal 默认字体
    normal = doc.styles['Normal']
    normal.font.name = DEFAULT_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(meta.get("title", "差旅费用统计单"))
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = DEFAULT_FONT
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # 元信息
    meta_para = doc.add_paragraph()
    meta_text = f"出行人：{meta['traveler']}    公司：{meta['company']}"
    if meta.get("start_date"):
        meta_text += f"\n行程：{meta['start_date']} ~ {meta['end_date']}    目的地：{meta.get('destination', '')}"
    elif meta.get("subtitle"):
        meta_text += f"\n{meta['subtitle']}"
    mr = meta_para.add_run(meta_text)
    mr.font.size = Pt(11)
    mr.font.name = DEFAULT_FONT
    mr._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # 发票明细表
    doc.add_paragraph()  # 空行
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "日期", "类型", "线路/备注", "金额(¥)"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)

    for idx, inv in enumerate(invoices, 1):
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(idx))
        _set_cell_text(cells[1], str(inv.date))
        _set_cell_text(cells[2], getattr(inv, "invoice_type", "") or "")
        _set_cell_text(cells[3], _route_str(inv))
        _set_cell_text(cells[4], f"{inv.amount:.2f}")

    # 合计
    totals = _category_totals(invoices)
    grand = sum(inv.amount for inv in invoices)
    doc.add_paragraph()
    total_para = doc.add_paragraph()
    tr = total_para.add_run("分类合计：\n")
    tr.bold = True
    tr.font.name = DEFAULT_FONT
    tr._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    detail = "  ".join(f"{k}：¥{v:.2f}" for k, v in totals.items() if v > 0)
    dr = total_para.add_run(detail + "\n")
    dr.font.name = DEFAULT_FONT
    dr._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    grand_run = total_para.add_run(f"\n总计：¥{grand:.2f}（大写：{_number_to_chinese(grand)}）")
    grand_run.bold = True
    grand_run.font.size = Pt(12)
    grand_run.font.name = DEFAULT_FONT
    grand_run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    out_path = folder / SUMMARY_FILENAME
    doc.save(out_path)
    logger.info(f"Generated summary sheet: {out_path}")
    return out_path


def generate_for_folder(folder: Path, kind: str = "trip") -> Optional[Path]:
    """扫描文件夹内 *.pdf，生成 统计单.docx。返回 docx 路径（无发票则返回 None）。"""
    folder = Path(folder)
    if not folder.is_dir():
        logger.warning(f"Not a directory: {folder}")
        return None

    invoices = collect_folder_invoices(folder)
    if not invoices:
        logger.info(f"No parseable invoices in {folder}, skipping summary sheet")
        return None

    meta = _meta_from_folder(folder, kind)
    return _build_docx(folder, invoices, meta)
